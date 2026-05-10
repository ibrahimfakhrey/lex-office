"""Extract text (or page images) from uploaded judgment files (PDF / DOCX).

Two paths:

  1. Digital PDF / DOCX → text-only ExtractResult, sent to Claude Haiku.
  2. Scanned PDF (no embedded text) → render pages to PNG images and
     return them in `images`. Routed to Claude Sonnet vision for OCR +
     structured extraction in one call.

The detection heuristic is intentionally simple: if the embedded text
density is below a small threshold, we assume the PDF is a scan and
fall back to vision. False positives (a digital PDF flagged as scanned)
just cost a bit more in vision tokens; false negatives (a scan slipping
through as text) are caught by Claude returning empty fields, which the
lawyer can correct manually.
"""
from __future__ import annotations

import io
import os
from dataclasses import dataclass, field
from typing import List


# Anything below this threshold on a non-trivial PDF (>=1 page) suggests the
# file is a scan with no embedded text. Real digital judgments easily clear
# 200 chars even on a single short page.
_MIN_USEFUL_CHARS_PER_PAGE = 80

# Some PDFs embed Arabic *visually* (using fonts with no toUnicode mapping)
# so PyMuPDF extracts characters from the Arabic Presentation Forms blocks
# (U+FB50–U+FDFF, U+FE70–U+FEFF) — the disconnected isolated/initial/medial/
# final glyphs — instead of base Arabic letters (U+0600–U+06FF). This text
# is unreadable for Claude (it's gibberish even to Arabic readers). We
# detect this and route to vision OCR, which can read the rendered page.
#
# Real digital Arabic PDFs have ~0% presentation-form chars; real-world
# broken encodings we've seen are 50-80%. 30% is a conservative threshold.
_MAX_PRESENTATION_FORM_RATIO = 0.30


def _is_presentation_form(c: str) -> bool:
    """True if char is in Arabic Presentation Forms-A or -B."""
    cp = ord(c)
    return (0xFB50 <= cp <= 0xFDFF) or (0xFE70 <= cp <= 0xFEFF)


def _is_base_arabic(c: str) -> bool:
    """True if char is a base Arabic letter (Arabic block + Supplement)."""
    cp = ord(c)
    return 0x0600 <= cp <= 0x06FF or 0x0750 <= cp <= 0x077F


def _has_broken_arabic_encoding(text: str) -> bool:
    """Heuristic: text has so many presentation-form glyphs that it's likely
    a PDF with broken Unicode mapping — should be re-OCR'd via vision.
    """
    pres = sum(1 for c in text if _is_presentation_form(c))
    base = sum(1 for c in text if _is_base_arabic(c))
    total = pres + base
    if total < 50:
        return False  # not enough Arabic to judge
    return (pres / total) > _MAX_PRESENTATION_FORM_RATIO

# Cap text we send to Claude. Most judgments are under ~30K chars; a hard cap
# keeps us under Haiku's input limits and bounds cost predictably.
MAX_CHARS = 60_000

# Vision OCR caps. Claude vision charges per image; capping pages bounds the
# cost of a single upload. 10 pages is enough for ~99% of court judgments
# and costs ~$0.05 with Sonnet 4.6 — still cheap relative to lawyer time.
MAX_VISION_PAGES = 10
# Render PDF pages at 2× zoom (≈ 144 DPI) — enough for accurate Arabic OCR
# without bloating image tokens. Higher DPI doesn't improve OCR meaningfully.
_VISION_RENDER_ZOOM = 2.0


class ExtractionError(Exception):
    """Generic extraction failure — message is in Arabic, safe to show user."""


class UnsupportedFileTypeError(ExtractionError):
    pass


# Kept for backwards compatibility (previous routes catch it). With OCR enabled
# this is no longer raised on scanned PDFs — they go through the vision path.
class ScannedPdfError(ExtractionError):
    pass


@dataclass
class ExtractResult:
    text: str
    page_count: int
    char_count: int
    source: str           # 'pdf' | 'docx'
    # 'text' for digital PDFs/DOCX; 'vision' for scanned PDFs that need OCR.
    mode: str = 'text'
    # PNG image bytes per page — only populated when mode='vision'.
    images: List[bytes] = field(default_factory=list)


# ── PDF ───────────────────────────────────────────────────────────────────────

def _render_pdf_pages_to_images(doc, max_pages: int = MAX_VISION_PAGES) -> List[bytes]:
    """Render up to `max_pages` PDF pages to PNG bytes for Claude vision."""
    import fitz  # PyMuPDF
    matrix = fitz.Matrix(_VISION_RENDER_ZOOM, _VISION_RENDER_ZOOM)
    images = []
    for i, page in enumerate(doc):
        if i >= max_pages:
            break
        pix = page.get_pixmap(matrix=matrix, alpha=False)
        images.append(pix.tobytes('png'))
    return images


def _extract_pdf(path: str) -> ExtractResult:
    import fitz  # PyMuPDF — best Arabic / RTL handling per 2026 evals

    doc = fitz.open(path)
    try:
        pages = []
        for page in doc:
            # "text" sort=True gives reading-order output, which on RTL Arabic
            # PDFs is what we actually want. Without sort, columns interleave.
            pages.append(page.get_text("text", sort=True))
        text = "\n\n".join(pages).strip()
        page_count = doc.page_count

        if page_count == 0:
            raise ExtractionError('الملف فارغ أو تالف — يرجى المحاولة بنسخة أخرى')

        # Three reasons to fall through to vision OCR:
        #  (a) Scanned PDF: no embedded text at all
        #  (b) Broken-Unicode PDF: text exists but uses presentation forms,
        #      not base Arabic letters — gibberish to Claude
        #  (c) Mixed/partial: too little usable text per page
        too_short = len(text) < _MIN_USEFUL_CHARS_PER_PAGE * page_count
        broken_arabic = _has_broken_arabic_encoding(text)
        if too_short or broken_arabic:
            images = _render_pdf_pages_to_images(doc)
            if not images:
                raise ExtractionError(
                    'تعذّر تحويل صفحات الملف إلى صور للتحليل — '
                    'يرجى المحاولة بنسخة أخرى'
                )
            return ExtractResult(
                text='', page_count=page_count, char_count=0,
                source='pdf', mode='vision', images=images,
            )
    finally:
        doc.close()

    return ExtractResult(
        text=text[:MAX_CHARS], page_count=page_count,
        char_count=len(text), source='pdf', mode='text',
    )


# ── DOCX ──────────────────────────────────────────────────────────────────────

def _extract_docx(path: str) -> ExtractResult:
    from docx import Document  # python-docx, already in requirements

    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Tables (used for case headers in many Arabic judgment templates)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    paragraphs.append(cell_text)

    text = "\n".join(paragraphs).strip()
    if not text:
        raise ExtractionError('الملف لا يحتوي على نص — يرجى المحاولة بنسخة أخرى')

    return ExtractResult(
        text=text[:MAX_CHARS], page_count=0,
        char_count=len(text), source='docx', mode='text',
    )


# ── Public API ────────────────────────────────────────────────────────────────

_PDF_EXTS = {'.pdf'}
_DOCX_EXTS = {'.docx'}
SUPPORTED_EXTS = _PDF_EXTS | _DOCX_EXTS


def extract_text(file_path: str) -> ExtractResult:
    """Detect type and extract. Raises ExtractionError subclasses on failure.

    All error messages are Arabic and safe to surface to the end user.
    Result.mode is 'text' for digital files; 'vision' for scanned PDFs
    where the caller should pass result.images to Claude vision instead.
    """
    ext = os.path.splitext(file_path)[1].lower()
    if ext in _PDF_EXTS:
        return _extract_pdf(file_path)
    if ext in _DOCX_EXTS:
        return _extract_docx(file_path)
    raise UnsupportedFileTypeError(
        'نوع الملف غير مدعوم — يرجى رفع PDF أو DOCX فقط'
    )
